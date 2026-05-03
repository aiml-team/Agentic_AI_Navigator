import io

from fastapi import HTTPException


def extract_pdf_text(content: bytes) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        parts  = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                pass
        return "\n".join(parts)
    except ImportError:
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                return "\n".join(
                    page.extract_text() or "" for page in pdf.pages
                )
        except ImportError:
            raise HTTPException(
                500,
                "PDF parsing library not installed. Run: pip install pypdf"
            )


def extract_docx_text(content: bytes) -> str:
    try:
        from docx import Document
        doc   = Document(io.BytesIO(content))
        parts = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except ImportError:
        raise HTTPException(
            500,
            "python-docx library not installed. Run: pip install python-docx"
        )
